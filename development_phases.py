# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_phases_doc():
    doc = Document()
    
    # 设置文档样式
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('中医体质分析系统开发阶段说明', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 需求分析
    doc.add_heading('1. 需求分析', 1)
    doc.add_paragraph('''本系统的功能需求分析如下：
1.1 核心功能需求
    • 用户可以选择多个症状
    • 系统根据症状匹配体质类型
    • 展示分析结果和养生建议
    
1.2 用户体验需求
    • 界面简洁直观
    • 操作流程简单
    • 结果展示清晰
    
1.3 技术需求
    • 支持多平台访问
    • 响应速度快
    • 易于维护和扩展''')

    # 2. 用户界面设计
    doc.add_heading('2. 用户界面设计', 1)
    doc.add_paragraph('''界面设计考虑以下要素：
2.1 布局设计
    • 采用卡片式布局
    • 居中对齐的内容
    • 清晰的视觉层次
    
2.2 交互元素
    • 多选框组件设计
    • 提交按钮设计
    • 结果展示区域设计
    
2.3 响应式适配
    • Bootstrap栅格系统
    • 移动端优化
    • 自适应布局''')

    # 3. 数据结构规划
    doc.add_heading('3. 数据结构规划', 1)
    doc.add_paragraph('''数据结构设计如下：
3.1 体质类型数据
    • 采用JSON对象存储
    • 包含症状、特征和建议
    • 便于扩展和维护
    
3.2 症状数据结构
    • 数组形式存储
    • 与体质类型关联
    • 支持多对多关系
    
3.3 分析结果结构
    • 包含体质类型
    • 包含特征描述
    • 包含养生建议''')

    # 4. 实现过程
    doc.add_heading('4. 实现过程', 1)
    doc.add_paragraph('''具体实现过程如下：
4.1 环境搭建
    • 安装Jekyll框架
    • 配置开发环境
    • 创建项目结构
    
4.2 页面开发
    • 实现HTML结构
    • 添加样式定义
    • 编写交互脚本
    
4.3 功能实现
    • 开发症状选择功能
    • 实现匹配算法
    • 完成结果展示''')

    doc.save('d:\\PV\\郭海静\\考研\\LINGNAN\\homework\\healthcare1\\开发阶段说明.docx')

if __name__ == '__main__':
    create_phases_doc()
