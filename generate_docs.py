# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_doc():
    doc = Document()
    
    # 设置文档样式
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('中医体质分析系统技术文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 项目概述
    doc.add_heading('1. 项目概述', 1)
    doc.add_paragraph('''本项目是一个基于Web技术的中医体质分析系统，通过分析用户的症状特征来判断其体质类型，并提供相应的养生建议。系统采用现代前端技术栈，结合传统中医理论，实现了一个简单但实用的健康咨询工具。''')

    # 2. 技术架构
    doc.add_heading('2. 技术架构', 1)
    doc.add_paragraph('''
• 前端框架：Jekyll静态站点生成器
• UI框架：Bootstrap 5
• 交互实现：原生JavaScript
• 数据存储：JSON对象
• 部署平台：GitHub Pages
''')

    # 3. 核心功能实现
    doc.add_heading('3. 核心功能实现', 1)
    doc.add_paragraph('''
体质分析算法的核心实现包括：
• 症状数据收集和预处理
• 多条件匹配算法
• 结果评分和权重计算
• 推荐系统逻辑
''')

    # 4. 数据结构设计
    doc.add_heading('4. 数据结构设计', 1)
    doc.add_paragraph('''
系统采用JSON对象存储体质类型数据，包含：
• 症状特征集合
• 体质类型描述
• 养生茶饮推荐
''')

    # 5. 界面设计
    doc.add_heading('5. 界面设计', 1)
    doc.add_paragraph('''
用户界面设计考虑以下几个方面：
• 清晰的视觉层次
• 响应式布局适配
• 良好的用户体验
• 中医特色元素
''')

    # 6. 部署说明
    doc.add_heading('6. 部署说明', 1)
    doc.add_paragraph('''
部署步骤包括：
1. 环境准备
   • 安装Jekyll
   • 配置依赖
   • 设置Git仓库

2. 构建过程
   • 生成静态文件
   • 资源优化
   • 部署验证

3. 维护更新
   • 定期检查
   • 内容更新
   • 性能优化
''')

    doc.save('d:\\PV\\郭海静\\考研\\LINGNAN\\homework\\healthcare1\\中医体质分析系统技术文档.docx')

if __name__ == '__main__':
    create_doc()
