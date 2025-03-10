# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_code_block(doc, code, filename, description=""):
    doc.add_heading(filename, level=2)
    if description:
        doc.add_paragraph(description)
    p = doc.add_paragraph()
    code_run = p.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)
    doc.add_paragraph()  # 添加空行

def create_doc():
    doc = Document()
    
    # 设置文档样式
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('中医体质分析系统项目代码文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Gemfile
    gemfile_content = '''source "https://rubygems.org"
gem "jekyll"
gem "webrick"
gem "minima"'''
    add_code_block(doc, gemfile_content, 'Gemfile', '项目依赖配置文件')

    # _config.yml
    config_content = '''title: 中医体质分析系统
description: 基于传统中医理论的体质分析工具
baseurl: ""
url: ""
theme: minima'''
    add_code_block(doc, config_content, '_config.yml', 'Jekyll配置文件')

    # default.html
    default_layout = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ page.title }}</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        .symptom-list {
            max-width: 600px;
            margin: 0 auto;
        }
        .card {
            border-radius: 10px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        .result-section {
            background-color: #f8f9fa;
            padding: 20px;
        }
    </style>
</head>
<body>
    {{ content }}
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>'''
    add_code_block(doc, default_layout, '_layouts/default.html', '默认布局模板')

    # index.html
    index_content = '''---
layout: default
title: 体质分析结果
---

<div class="container">
    <div class="text-center my-4">
        <h2>您的体质分析报告</h2>
        
        <!-- 选择症状部分 -->
        <div class="card mb-4">
            <div class="card-body">
                <h4>请选择您的主要症状（可多选）：</h4>
                <div class="symptom-list mt-3">
                    <div class="form-check mb-2">
                        <input type="checkbox" class="form-check-input" id="s1" value="失眠多梦">
                        <label class="form-check-label" for="s1">失眠多梦</label>
                    </div>
                    <div class="form-check mb-2">
                        <input type="checkbox" class="form-check-input" id="s2" value="手脚发热">
                        <label class="form-check-label" for="s2">手脚心发热</label>
                    </div>
                    <div class="form-check mb-2">
                        <input type="checkbox" class="form-check-input" id="s3" value="容易疲劳">
                        <label class="form-check-label" for="s3">容易疲劳</label>
                    </div>
                    <div class="form-check mb-2">
                        <input type="checkbox" class="form-check-input" id="s4" value="怕冷">
                        <label class="form-check-label" for="s4">怕冷手脚凉</label>
                    </div>
                    <div class="form-check mb-2">
                        <input type="checkbox" class="form-check-input" id="s5" value="消化不良">
                        <label class="form-check-label" for="s5">消化不良</label>
                    </div>
                    <button class="btn btn-primary mt-3" onclick="analyzeConstitution()">生成分析结果</button>
                </div>
            </div>
        </div>

        <!-- 分析结果部分 -->
        <div id="result" style="display: none;" class="card">
            <div class="card-body">
                <h3 id="constitution-type" class="mb-3"></h3>
                <div id="characteristics" class="mb-3"></div>
                <div id="recommendations" class="mb-4"></div>
                <div class="text-muted small">
                    <p>注：本结果仅供参考，如有疑问请咨询专业医师</p>
                </div>
            </div>
        </div>
    </div>
</div>

<script>
const constitutionData = {
    '阴虚质': {
        symptoms: ['失眠多梦', '手脚发热'],
        characteristics: '体质偏热，容易口干、手脚心发热、睡眠质量差',
        recommendations: ['菊花枸杞茶', '石斛养生茶', '桑菊茶']
    },
    '阳虚质': {
        symptoms: ['怕冷', '手脚凉'],
        characteristics: '体质偏寒，容易怕冷、手脚发凉、精神疲惫',
        recommendations: ['生姜红糖茶', '桂圆红枣茶', '艾草红糖茶']
    },
    '气虚质': {
        symptoms: ['容易疲劳', '消化不良'],
        characteristics: '体质虚弱，容易疲劳、气短乏力、食欲不振',
        recommendations: ['人参养生茶', '黄芪当归茶', '西洋参茶']
    }
};

function analyzeConstitution() {
    const selectedSymptoms = Array.from(document.querySelectorAll('input[type="checkbox"]:checked'))
        .map(cb => cb.value);
    
    let matchedType = null;
    let maxMatches = 0;

    for (const [type, data] of Object.entries(constitutionData)) {
        const matches = data.symptoms.filter(s => selectedSymptoms.includes(s)).length;
        if (matches > maxMatches) {
            maxMatches = matches;
            matchedType = type;
        }
    }

    if (matchedType) {
        const data = constitutionData[matchedType];
        document.getElementById('constitution-type').textContent = `您的体质类型：${matchedType}`;
        document.getElementById('characteristics').innerHTML = `
            <p class="lead">${data.characteristics}</p>
        `;
        document.getElementById('recommendations').innerHTML = `
            <h4>推荐养生茶饮</h4>
            <ul class="list-unstyled">
                ${data.recommendations.map(tea => `<li>• ${tea}</li>`).join('')}
            </ul>
        `;
        document.getElementById('result').style.display = 'block';
        document.getElementById('result').scrollIntoView({ behavior: 'smooth' });
    }
}
</script>'''
    add_code_block(doc, index_content, 'index.html', '主页面文件')

    # .gitignore
    gitignore_content = '''_site
.sass-cache
.jekyll-cache
.jekyll-metadata
vendor'''
    add_code_block(doc, gitignore_content, '.gitignore', 'Git忽略文件配置')

    doc.save('d:\\PV\\郭海静\\考研\\LINGNAN\\homework\\healthcare1\\项目代码文档.docx')

if __name__ == '__main__':
    create_doc()
